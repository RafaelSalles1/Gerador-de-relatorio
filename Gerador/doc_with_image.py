from docx import Document
from docx.shared import Inches
import tkinter as tk
from tkinter import filedialog

root = tk.Tk()
document = Document()
document.add_heading("Titulo", 0)
canvas1 = tk.Canvas(root, width = 300, height = 300, bg = 'lightsteelblue')
canvas1.pack()

def RequestPhoto():
    global lst
    #filename = filedialog.askopenfilename()
    filez = filedialog.askopenfilenames()
    #print(filez)
    lst = list(filez)
    print(lst)
    root.destroy()
    return lst


search_button = tk.Button(text="Insira foto", command = RequestPhoto, bg = 'lightgreen', fg = 'white', font = ('helvetica', 12, 'bold'))
canvas1.create_window(150, 150, window = search_button)


root.mainloop()
for filex in lst:
    document.add_picture(filex, width=Inches(1.25))
#root.destroy()
document.save('demo2.docx')
