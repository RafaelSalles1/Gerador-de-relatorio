from docx import Document
from docx.shared import Inches
import tkinter as tk
from tkinter import filedialog
import pandas as pd


root = tk.Tk()

canvas1 = tk.Canvas(root, width = 300, height = 300, bg = 'lightsteelblue')
canvas1.pack()

document = Document()

document.add_heading('Document Title', 0)

p = document.add_paragraph("A plain paragraph having some ")
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic').italic = True

document.add_heading('Heading, level 1', level = 1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)

document.add_paragraph(
    'first item in ordered list', style='List Number'
)
def getPhoto():


    filename = filedialog.askopenfilename()
    return filename

browseButton_Photo = tk.Button(text='Import Image', command=getPhoto, bg='green', fg='white', font=('helvetica', 12, 'bold'))
canvas1.create_window(150, 150, window=browseButton_Photo)


records = (
    (3, "101", "Spam"),
    (7, "422", "Eggs"),
    (4, "631", 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdrs_cells = table.rows[0].cells
hdrs_cells[0].text = 'Qtdy'
hdrs_cells[1].text = 'ID'
hdrs_cells[2].text = 'Desc'
for qty, id, desc in records:
    rows_cells = table.add_row().cells
    rows_cells[0].text = str(qty)
    rows_cells[1].text = str(id)
    rows_cells[2].text = str(desc)
root.mainloop()
#document.add_picture(filename, width = Inches(2.5))
document.add_page_break()
document.save('demo.docx')
