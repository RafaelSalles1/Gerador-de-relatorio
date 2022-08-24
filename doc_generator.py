from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.section import *
from docx.enum.text import *
from docx.enum.style import *
from docx.enum.dml import *
import tkinter as tk
from tkinter import filedialog
import pandas as pd

document = Document()

#SETANDO TEXTO NORMAL COMO "CENTRALIZADO"
stilo = document.styles['Normal']
formatacao = stilo.paragraph_format
formatacao.alignment = WD_ALIGN_PARAGRAPH.CENTER

#CRIANDO UM ESTILO ESPECIFICO PARA O TITULO E EDITANDO SUAS CARACTERISTICAS
style = document.styles.add_style('Titulo', WD_STYLE_TYPE.PARAGRAPH)
para_format = style.paragraph_format
para_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_format.line_spacing = 1.5
para_format.right_indent = Cm(1.09)
para_format.line_spacing = Cm(1.50)
para_format.space_after = 0
font = style.font
font.bold = True
font.name = "Verdana"
font.size = Pt(36)
font.all_caps = True

#CRIANDO UM ESTILO ESPECIFICO PARA O SUBTITULOTITULO E EDITANDO SUAS CARACTERISTICAS
style = document.styles.add_style('Subtitulo', WD_STYLE_TYPE.PARAGRAPH)
para_format = style.paragraph_format
para_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_format.right_indent = Cm(-0.68)
para_format.line_spacing = Cm(0)
para_format.space_after = 0
font = style.font
font.bold = True
font.name = "Verdana"
font.size = Pt(12)
font.highlight_color = WD_COLOR_INDEX.GRAY_50

#CRIANDO UM ESTILO ESPECIFICO PARA O TEXTO EM GERAL E EDITANDO SUAS CARACTERISTICAS
style = document.styles.add_style('texto', WD_STYLE_TYPE.PARAGRAPH)
para_format = style.paragraph_format
para_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
para_format.right_indent = Cm(0)
para_format.line_spacing = Cm(0)
para_format.space_after = 0
font = style.font
font.name = "Verdana"
font.size = Pt(12)

#ABRIR O EXCEL
import_file_path = r"X:\SERVIÇOS\TM-S\Relatório Padrão TMS - Não Apagar.xlsx"
df = pd.read_excel(import_file_path)

#DEIXAR A FOLHA EM PAISAGEM E COM AS MARGENS CORRETAS
section = document.sections[-1]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = Cm(29.7)
section.page_height = Cm(21.0)
section.bottom_margin = Cm(3.00)
section.left_margin = Cm(1.30)
section.right_margin = Cm(3.20)
section.top_margin = Cm(2.60)

#ANEXANDO O CONTEUDO DO EXCEL AS DEVIDAS VARIAVEIS
titulo = "\tRELATÓRIO DE MELHORIAS NO SISTEMA DE VAPOR E DRENAGEM DE CONDENSADO"

subtitulo_introducao = "\t" + df["TÍTULO"][0]
texto_introducao = "\t" + df["TEXTO"][0]

subtitulo_melhorias = "\t" + df["TÍTULO"][1]
texto_melhorias = "\t" + df["TEXTO"][1]
imagem_melhorias = r"X:\SERVIÇOS\TM-S\imagens_relatorio\imagem_melhorias.jpg"

subtitulo_geracao = "\t" + df['TÍTULO'][2]
legenda_geracao = "\t" + df["Figuras"][2] 
tabela_geracao = r"X:\SERVIÇOS\TM-S\imagens_relatorio\tabela_geracao.jpg"

subtitulo_distribuicao = "\t" + df['TÍTULO'][3]
texto_distribuicao = "\t" + df["TEXTO"][3]
imagem_distribuicao = r"X:\SERVIÇOS\TM-S\imagens_relatorio\imagem_distribuicao.jpg"

subtitulo_boaspraticas = "\t" + df['TÍTULO'][4]
texto_boaspraticas = "\t" + df["TEXTO"][4]
texto2_boaspraticas = "\t" + df["TEXTO"][5]
imagem_boaspraticas = r"X:\SERVIÇOS\TM-S\imagens_relatorio\imagem_boaspraticas.jpg"
imagem2_boaspraticas = r"X:\SERVIÇOS\TM-S\imagens_relatorio\imagem2_boaspraticas.jpg"

subtitulo_antesred = "\t" + df['TÍTULO'][6]
texto_antesred = "\t" + df["TEXTO"][6]
imagem_antesred = r"X:\SERVIÇOS\TM-S\imagens_relatorio\imagem_antesred.jpg"
imagem2_antesred = r"X:\SERVIÇOS\TM-S\imagens_relatorio\imagem2_antesred.jpg"

subtitulo_adequacao = "\t" + df['TÍTULO'][7]
texto_adequacao = "\t" + df["TEXTO"][7]
imagem_adequacao = r"X:\SERVIÇOS\TM-S\imagens_relatorio\imagem_adequacao.jpg"
imagem2_adequacao = r"X:\SERVIÇOS\TM-S\imagens_relatorio\imagem2_adequacao.jpg"
tabela_adequacao = r"X:\SERVIÇOS\TM-S\imagens_relatorio\tabela_adequacao.jpg"
problema_adequacao = "\t" + df["TEXTO4"][7]

subtitulo_desaeracao = "\t" + df['TÍTULO'][8]
texto_desaeracao = "\t" + df["TEXTO"][8]
imagem_desaeracao = r"X:\SERVIÇOS\TM-S\imagens_relatorio\imagem_desaeracao.jpg"
imagem2_desaeracao = r"X:\SERVIÇOS\TM-S\imagens_relatorio\imagem2_desaeracao.jpg"
nota_desaeracao = "\t" + df["TEXTO2"][8]

subtitulo_drenagem_coletiva = "\t" + df['TÍTULO'][9]
texto_drenagem_coletiva = "\t" + df["TEXTO"][9]
texto2_drenagem_coletiva = "\t" + df["Figuras"][9]
imagem_drenagem_coletiva = r"X:\SERVIÇOS\TM-S\imagens_relatorio\imagem_drenagem_coletiva.jpg"
texto3_drenagem_coletiva = "\t" + df["TEXTO2"][9]

subtitulo_vapor_preso = "\t" + df['TÍTULO'][10]
texto_vapor_preso = "\t" + df["TEXTO"][10]
texto2_vapor_preso = "\t" + df["Figuras"][10]
imagem_vapor_preso = r"X:\SERVIÇOS\TM-S\imagens_relatorio\imagem_vapor_preso.jpg"
texto3_vapor_preso = "\t" + df["TEXTO2"][10]

subtitulo_vapor_preso_cilindros = "\t" + df['TÍTULO'][11]
texto_vapor_preso_cilindros = "\t" + df["TEXTO"][11]
imagem_vapor_preso_cilindros = r"X:\SERVIÇOS\TM-S\imagens_relatorio\imagem_vapor_preso_cilindros.jpg"
texto2_vapor_preso_cilindros = "\t" + df["TEXTO1"][11]
imagem2_vapor_preso_cilindros = r"X:\SERVIÇOS\TM-S\imagens_relatorio\imagem2_vapor_preso_cilindros.jpg"
texto3_vapor_preso_cilindros = "\t" + df["TEXTO3"][11]
imagem3_vapor_preso_cilindros = r"X:\SERVIÇOS\TM-S\imagens_relatorio\imagem3_vapor_preso_cilindros.jpg"
texto4_vapor_preso_cilindros = "\t" + df["TEXTO5"][11]
texto5_vapor_preso_cilindros = "\t" + df["TEXTO6"][11]
imagem4_vapor_preso_cilindros = r"X:\SERVIÇOS\TM-S\imagens_relatorio\imagem4_vapor_preso_cilindros.jpg"

subtitulo_isolamento = "\t" + df['TÍTULO'][12]
texto_isolamento = "\t" + df["TEXTO"][12]
tabela_isolamento = r"X:\SERVIÇOS\TM-S\imagens_relatorio\tabela_isolamento.jpg"
tabela2_isolamento = r"X:\SERVIÇOS\TM-S\imagens_relatorio\tabela2_isolamento.jpg"

subtitulo_boia_termo = "\t" + df['TÍTULO'][13]
texto_boia_termo = "\t" + df["TEXTO"][13]
grafico_boia_termo = r"X:\SERVIÇOS\TM-S\imagens_relatorio\grafico_boia_termo.jpg"
texto2_boia_termo = "\t" + df["TEXTO1"][13]

subtitulo_vaz_ext = "\t" + df['TÍTULO'][14]
texto_vaz_ext = "\t" + df["TEXTO"][14]
tabela_vaz_ext = r"X:\SERVIÇOS\TM-S\imagens_relatorio\tabela_vaz_ext.jpg"

subtitulo_bypass = "\t" + df['TÍTULO'][15]
texto_bypass = "\t" + df["TEXTO"][15]
imagem_bypass = r"X:\SERVIÇOS\TM-S\imagens_relatorio\imagem_bypass.jpg"

#ADICIONANDO O TITULO NO DOCUMENTO
document.add_paragraph('', style = 'Titulo')
document.add_paragraph('', style = 'Titulo')
document.add_paragraph('', style = 'Titulo')
titulo_add = document.add_paragraph(titulo, style = 'Titulo')
titulo_add.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
document.add_page_break()

#ADICIONANDO INTRODUCAO
document.add_paragraph(subtitulo_introducao, style = 'Subtitulo')
document.add_paragraph()
paragrafo_introducao = document.add_paragraph(texto_introducao, style = 'texto')
paragrafo_introducao.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
document.add_page_break()

#ADICIONANDO MELHORIAS
document.add_paragraph(subtitulo_melhorias, style = 'Subtitulo')
document.add_paragraph()
paragrafo_melhorias = document.add_paragraph(texto_melhorias, style = 'texto')
paragrafo_melhorias.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
document.add_picture(imagem_melhorias, width = Cm(15.13))

#ADICIONANDO GERACAO DE VAPOR
document.add_paragraph(subtitulo_geracao, style = 'Subtitulo')
document.add_picture(tabela_geracao)
document.add_paragraph(legenda_geracao, style = "texto")
document.add_page_break()

#ADICIONANDO DISTRIBUICAO
document.add_paragraph(subtitulo_distribuicao, style = 'Subtitulo')
document.add_paragraph(texto_distribuicao, style = "texto")
document.add_picture(imagem_distribuicao)
document.add_page_break()

#ADICIONANDO BOAS PRATICAS
document.add_paragraph(subtitulo_boaspraticas, style = 'Subtitulo')
document.add_paragraph(texto_boaspraticas, style = "texto")
document.add_paragraph(texto2_boaspraticas, style = "texto")
document.add_picture(imagem_boaspraticas)
document.add_picture(imagem2_boaspraticas)

#ADICIONANDO NECESSIDADE DE DRENAGEM ANTES DE REDUCAO / CONTROLE
document.add_paragraph(subtitulo_antesred, style = 'Subtitulo')
document.add_paragraph(texto_antesred, style = "texto")
document.add_picture(imagem_antesred)
document.add_picture(imagem2_antesred)
document.add_page_break()

#ADICIONANDO ADEQUACAO DE LINHAS EXISTENTES
document.add_paragraph(subtitulo_adequacao, style = 'Subtitulo')
document.add_paragraph(texto_adequacao, style = "texto")
document.add_picture(imagem_adequacao)
document.add_picture(imagem2_adequacao)
document.add_picture(tabela_adequacao)
document.add_paragraph(problema_adequacao)
document.add_page_break()

#ADICIONANDO DESAERACAO
document.add_paragraph(subtitulo_desaeracao, style = 'Subtitulo')
document.add_paragraph(texto_desaeracao, style = "texto")
document.add_picture(imagem_desaeracao)
document.add_picture(imagem2_desaeracao)
document.add_paragraph(nota_desaeracao)
document.add_page_break()

#ADICIONANDO DRENAGEM COLETIVA
document.add_paragraph(subtitulo_drenagem_coletiva, style = 'Subtitulo')
document.add_paragraph(texto_drenagem_coletiva, style = "texto")
document.add_paragraph(texto2_drenagem_coletiva, style = "texto")
document.add_picture(imagem_drenagem_coletiva)
document.add_paragraph(texto3_drenagem_coletiva, style = "texto")
document.add_page_break()

#ADICIONANDO VAPOR PRESO TEORIA
document.add_paragraph(subtitulo_vapor_preso, style = 'Subtitulo')
document.add_paragraph(texto_vapor_preso, style = "texto")
document.add_paragraph(texto2_vapor_preso, style = "texto")
document.add_picture(imagem_vapor_preso)
document.add_paragraph(texto3_vapor_preso, style = "texto")
document.add_page_break()

#ADICIONANDO VAPOR PRESO CILINDROS
document.add_paragraph(subtitulo_vapor_preso_cilindros, style = 'Subtitulo')
document.add_paragraph(texto_vapor_preso_cilindros, style = "texto")
document.add_picture(imagem_vapor_preso_cilindros)
document.add_paragraph(texto2_vapor_preso_cilindros, style = "texto")
document.add_picture(imagem2_vapor_preso_cilindros)
document.add_paragraph(texto3_vapor_preso_cilindros, style = "texto")
document.add_picture(imagem3_vapor_preso_cilindros)
document.add_paragraph(texto4_vapor_preso_cilindros, style = "texto")
document.add_paragraph(texto5_vapor_preso_cilindros, style = "texto")
document.add_picture(imagem4_vapor_preso_cilindros)
document.add_page_break()

#ADICIONANDO ISOLAMENTO
document.add_paragraph(subtitulo_isolamento, style = 'Subtitulo')
document.add_paragraph(texto_isolamento, style = "texto")
document.add_picture(tabela_isolamento)
document.add_picture(tabela2_isolamento)
document.add_page_break()

#ADICIONANDO BOIA TERMO
document.add_paragraph(subtitulo_boia_termo, style = 'Subtitulo')
document.add_paragraph(texto_boia_termo, style = "texto")
document.add_picture(grafico_boia_termo)
document.add_paragraph(texto2_boia_termo, style = 'texto')
document.add_page_break()

#ADICIONANDO VAZAMENTO EXTERNO
document.add_paragraph(subtitulo_vaz_ext, style = 'Subtitulo')
document.add_paragraph(texto_vaz_ext, style = "texto")
document.add_picture(tabela_vaz_ext)
document.add_page_break()

#ADICIONANDO BYPASS
document.add_paragraph(subtitulo_bypass, style = 'Subtitulo')
document.add_paragraph(texto_bypass, style = "texto")
document.add_picture(imagem_bypass)
document.add_page_break()



document.save('foobar.docx')





