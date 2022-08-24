from tkinter import *
from tkinter import PhotoImage, messagebox
import tkinter as tk
import PIL
from tkinter import filedialog
from PIL import Image
from PIL import ImageTk
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.section import *
from docx.enum.text import *
from docx.enum.style import *
from docx.enum.dml import *
import pandas as pd


class MainApp():
    def __init__(self, root):

        self.root = root

        self.imagem_disparco = Image.open(r'C:\Users\rafael\Desktop\DESKTOP\rafael\projects\Gerador de relatorio\LOGO DISPARCO\PNG-transparente.png')
        self.imagem_disparco = self.imagem_disparco.resize((280, 130))
        self.logo_disparco = ImageTk.PhotoImage(self.imagem_disparco)

        self.canvas2 = tk.Canvas(root, width=900, height=501, bg="#D8D0C8", confine=TRUE)
        self.canvas2.pack()
        self.canvas2.create_image(150, 100, image=self.logo_disparco)

        self.num_proc = tk.Entry(self.canvas2)
        self.num_proc.place(width=100, height=30, x=700, y=100)

        self.nome_cliente = tk.Entry(self.canvas2)
        self.nome_cliente.place(width=300, height=30, x=500, y=50)

        self.frame1 = tk.Frame(self.canvas2, bg="#D8D0C8")
        self.frame1.place(relwidth=0.33, relheight=0.66, x=0, y=167)

        self.frame2 = tk.Frame(self.canvas2, bg="#D8D0C8")
        self.frame2.place(relwidth=0.33, relheight=0.66, x=301, y=167)

        self.frame3 = tk.Frame(self.canvas2, bg="#D8D0C8")
        self.frame3.place(relwidth=0.33, relheight=0.66, x=601, y=167)

        self.label_nome_cliente = tk.Label(self.canvas2, text="Nome do cliente:", bg="#D8D0C8", fg="black",
                                           font=("helvetica", 9, "bold"))
        self.label_nome_cliente.place(width=100, height=30, x=400, y=50)

        self.label_num_proc = tk.Label(self.canvas2, text="Numero do processo:", bg="#D8D0C8", fg="black",
                                       font=("helvetica", 9, 'bold'))
        self.label_num_proc.place(width=128, height=30, x=400, y=100)

        self.generate_button = tk.Button(self.frame3, text="Gerar Relatório", bg="#D8D0C8", fg="black",
                                         command=self.click)
        self.generate_button.place(relwidth=0.5, relheight=0.1, x=120, y=270)

        self.var_introducao = tk.BooleanVar()
        self.var_melhorias = tk.IntVar()
        self.var_geracao_de_vapor = tk.IntVar()
        self.var_distribuicao_de_vapor = tk.IntVar()
        self.var_boas_praticas = tk.IntVar()
        self.var_adeq_in_out = tk.IntVar()
        self.var_adeq_linhas_ext = tk.IntVar()
        self.var_desaer_rede = tk.IntVar()
        self.var_dren_coletiva = tk.IntVar()
        self.var_vapor_preso_teoria = tk.IntVar()
        self.var_vapor_preso_cilindros = tk.IntVar()
        self.var_tubos_isolamento = tk.IntVar()
        self.var_purgadores_termo_x_boia = tk.IntVar()
        self.var_locais_vazamentos_ext = tk.IntVar()
        self.var_valvulas_bypass = tk.IntVar()

        self.cb_introducao = tk.Checkbutton(self.frame1, text="Introdução", bg="#D8D0C8", variable=self.var_introducao,
                                            pady=5, padx=0)
        self.cb_introducao.place(relx=0, rely=0)

        self.cb_melhorias = tk.Checkbutton(self.frame1, text="Melhorias", bg="#D8D0C8", variable=self.var_melhorias,
                                           pady=5)
        self.cb_melhorias.place(relx=0, rely=0.15)

        self.cb_geracao_de_vapor = tk.Checkbutton(self.frame1, text="Geração de Vapor (teoria)", bg="#D8D0C8",
                                                  variable=self.var_geracao_de_vapor, pady=5)
        self.cb_geracao_de_vapor.place(relx=0, rely=0.30)

        self.cb_distribuicao_de_vapor = tk.Checkbutton(self.frame1, text="Distribuição de Vapor", bg="#D8D0C8",
                                                       variable=self.var_distribuicao_de_vapor, pady=5,
                                                       command=self.check_status_distribuicao)
        self.cb_distribuicao_de_vapor.place(relx=0, rely=0.45)

        self.cb_boas_praticas = tk.Checkbutton(self.frame1, text="Boas praticas para remoção de condensado",
                                               bg="#D8D0C8",
                                               variable=self.var_boas_praticas, pady=5,
                                               command=self.check_status_boas_praticas)
        self.cb_boas_praticas.place(relx=0, rely=0.60)

        self.cb_adeq_in_out = tk.Checkbutton(self.frame2, text="Adequação necessidade dren. Inlet/Outlet", bg="#D8D0C8",
                                             variable=self.var_adeq_in_out, pady=5, command=self.check_status_antes_red)
        self.cb_adeq_in_out.place(relx=0, rely=0)

        self.cb_adeq_linhas_ext = tk.Checkbutton(self.frame2, text="Adequação drenagens de linha existentes",
                                                 bg="#D8D0C8",
                                                 variable=self.var_adeq_linhas_ext, pady=5,
                                                 command=self.check_status_adeq_existente)
        self.cb_adeq_linhas_ext.place(relx=0, rely=0.15)

        self.cb_desaer_rede = tk.Checkbutton(self.frame2, text="Desaeração da rede de vapor", bg="#D8D0C8",
                                             variable=self.var_desaer_rede, pady=5,
                                             command=self.check_status_desaeracao)
        self.cb_desaer_rede.place(relx=0, rely=0.3)

        self.cb_dren_coletiva = tk.Checkbutton(self.frame2, text="Drenagem coletiva", bg="#D8D0C8",
                                               variable=self.var_dren_coletiva, pady=5,
                                               command=self.check_status_drenagem_coletiva)
        self.cb_dren_coletiva.place(relx=0, rely=0.45)

        self.cb_vapor_preso_teoria = tk.Checkbutton(self.frame2, text="Vapor Preso (Teoria)", bg="#D8D0C8",
                                                    variable=self.var_vapor_preso_teoria, pady=5,
                                                    command=self.check_status_vapor_preso)
        self.cb_vapor_preso_teoria.place(relx=0, rely=0.6)

        self.cb_vapor_preso_cilindros = tk.Checkbutton(self.frame3, text="Vapor Preso (Cilindros)", bg="#D8D0C8",
                                                       variable=self.var_vapor_preso_cilindros, pady=5,
                                                       command=self.check_status_vapor_preso_cilindros)
        self.cb_vapor_preso_cilindros.place(relx=0, rely=0)

        self.cb_tubos_isolamento = tk.Checkbutton(self.frame3, text="Tubulações sem isolamento", bg="#D8D0C8",
                                                  variable=self.var_tubos_isolamento, pady=5,
                                                  command=self.check_status_tubos_isolamento)
        self.cb_tubos_isolamento.place(relx=0, rely=0.15)

        self.cb_purgadores_termo_x_boia = tk.Checkbutton(self.frame3, text="Purgadores Termo x Boias", bg="#D8D0C8",
                                                         variable=self.var_purgadores_termo_x_boia, pady=5)
        self.cb_purgadores_termo_x_boia.place(relx=0, rely=0.3)

        self.cb_locais_vazamentos_ext = tk.Checkbutton(self.frame3, text="Locais com vaz. externos", bg="#D8D0C8",
                                                       variable=self.var_locais_vazamentos_ext, pady=5,
                                                       command=self.check_status_locais_vaz_ext)
        self.cb_locais_vazamentos_ext.place(relx=0, rely=0.45)

        self.cb_valvulas_bypass = tk.Checkbutton(self.frame3, text="Valvulas By-Pass", bg="#D8D0C8",
                                                 variable=self.var_valvulas_bypass, pady=5,
                                                 command=self.check_status_valvulas_bypass)
        self.cb_valvulas_bypass.place(relx=0, rely=0.6)

    def click(self):
        self.nome_cliente_str = self.nome_cliente.get()
        self.num_proc_str = self.num_proc.get()
        self.save_name = "RG-" + self.num_proc_str + "-001-20-00-Relatório de Gerenciamento-" + self.nome_cliente_str + ".docx"
        self.rota_imagens = 'C:/Users/rafael/Desktop/DESKTOP/rafael/projects/Gerador de relatorio/imagens_relatorio'
        document = Document()

        # CRIANDO UM ESTILO ESPECIFICO PARA O TITULO E EDITANDO SUAS CARACTERISTICAS
        style = document.styles.add_style('Titulo', WD_STYLE_TYPE.PARAGRAPH)
        para_format = style.paragraph_format
        #para_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_format.line_spacing = 1.5
        para_format.right_indent = Cm(1.09)
        para_format.line_spacing = Cm(1.50)
        para_format.space_after = 0
        font = style.font
        font.bold = True
        font.name = "Verdana"
        font.size = Pt(36)
        font.all_caps = True

        # CRIANDO UM ESTILO ESPECIFICO PARA O SUBTITULOTITULO E EDITANDO SUAS CARACTERISTICAS
        style = document.styles.add_style('Subtitulo', WD_STYLE_TYPE.PARAGRAPH)
        para_format = style.paragraph_format
        para_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para_format.right_indent = Cm(-0.68)
        para_format.line_spacing = Cm(0)
        para_format.space_after = 0
        font = style.font
        font.bold = True
        font.name = "Verdana"
        font.size = Pt(12)
        #		font.highlight_color = WD_COLOR_INDEX.GRAY_50

        # CRIANDO UM ESTILO ESPECIFICO PARA O TEXTO EM GERAL E EDITANDO SUAS CARACTERISTICAS
        style = document.styles.add_style('texto', WD_STYLE_TYPE.PARAGRAPH)
        para_format = style.paragraph_format
        #para_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para_format.right_indent = Cm(0)
        para_format.line_spacing = Cm(0)
        para_format.space_after = 0
        font = style.font
        font.name = "Verdana"
        font.size = Pt(12)

        # ABRIR O EXCEL
        import_file_path = r'C:\Users\rafael\Desktop\DESKTOP\rafael\projects\Gerador de relatorio\Relatório Padrão TMS - Não Apagar.xlsx'
        df = pd.read_excel(import_file_path)

        # DEIXAR A FOLHA EM PAISAGEM E COM AS MARGENS CORRETAS
        section = document.sections[-1]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)
        section.page_height = Cm(21.0)
        section.bottom_margin = Cm(3.00)
        section.left_margin = Cm(1.30)
        section.right_margin = Cm(3.20)
        section.top_margin = Cm(2.60)

        # ANEXANDO O CONTEUDO DO EXCEL AS DEVIDAS VARIAVEIS
        titulo = "\tRELATÓRIO DE MELHORIAS NO SISTEMA DE VAPOR E DRENAGEM DE CONDENSADO"

        subtitulo_introducao = "1." + "\t" + df["TÍTULO"][0]
        texto_introducao = "\t" + df["TEXTO"][0]
        texto2_introducao = "\t" + df["Figuras"][0]

        subtitulo_melhorias = "2." + "\t" + df["TÍTULO"][1]
        texto_melhorias = "\t" + df["TEXTO"][1]
        imagem_melhorias = self.rota_imagens + "/imagem_melhorias.jpg"

        subtitulo_geracao = "2.1." + "\t" + df['TÍTULO'][2]
        legenda_geracao = "\t" + df["Figuras"][2]
        texto_geracao = '\t' + df['TEXTO'][2]
        texto2_geracao = '\t' + df['TEXTO1'][2]
        tabela_geracao = self.rota_imagens + r"\tabela_geracao.jpg"

        subtitulo_distribuicao = "2.2." + "\t" + df['TÍTULO'][3]
        texto_distribuicao = "\t" + df["TEXTO"][3]
        texto2_distribuicao = "\t" + df["Figuras"][3]
        texto3_distribuicao = "\t" + df["TEXTO1"][3]
        texto4_distribuicao = "\t" + df["TEXTO2"][3]
        texto5_distribuicao = "\t" + df["TEXTO3"][3]
        texto6_distribuicao = "\t" + df["TEXTO4"][3]
        texto7_distribuicao = "\t" + df["TEXTO5"][3]
        imagem_distribuicao = self.rota_imagens + "\imagem_distribuicao.jpg"

        subtitulo_boaspraticas = "2.2.1." + "\t" + df['TÍTULO'][4]
        texto_boaspraticas = "\t" + df["TEXTO"][4]
        texto2_boaspraticas = df["Figuras"][4]
        texto3_boaspraticas = df["TEXTO1"][4]
        texto4_boaspraticas = df["TEXTO2"][4]
        texto5_boaspraticas = df["TEXTO3"][4]
        texto6_boaspraticas = "\t" + df["TEXTO"][5]
        imagem_boaspraticas = self.rota_imagens + r"\imagem_boaspraticas.jpg"
        imagem2_boaspraticas = self.rota_imagens + r"\imagem2_boaspraticas.jpg"

        subtitulo_antesred = "2.2.2." + "\t" + df['TÍTULO'][6]
        texto_antesred = "\t" + df["TEXTO"][6]
        imagem_antesred = self.rota_imagens + r"\imagem_antesred.jpg"
        imagem2_antesred = self.rota_imagens + r"\imagem2_antesred.jpg"

        subtitulo_adequacao = "2.2.3." + "\t" + df['TÍTULO'][7]
        texto_adequacao = "\t" + df["TEXTO"][7]
        imagem_adequacao = self.rota_imagens + r"\imagem_adequacao.jpg"
        imagem2_adequacao = self.rota_imagens + r"\imagem2_adequacao.jpg"
        tabela_adequacao = self.rota_imagens + r"\tabela_adequacao.jpg"
        problema_adequacao = "\t" + df["TEXTO4"][7]

        subtitulo_desaeracao = "2.2.4." + "\t" + df['TÍTULO'][8]
        texto_desaeracao = "\t" + df["TEXTO"][8]
        imagem_desaeracao = self.rota_imagens + r"\imagem_desaeracao.jpg"
        imagem2_desaeracao = self.rota_imagens + r"\imagem2_desaeracao.jpg"
        nota_desaeracao = "\t" + df["TEXTO2"][8]

        subtitulo_drenagem_coletiva = "2.3." + "\t" + df['TÍTULO'][9]
        texto_drenagem_coletiva = "\t" + df["TEXTO"][9]
        texto2_drenagem_coletiva = "\t" + df["Figuras"][9]
        imagem_drenagem_coletiva = self.rota_imagens + r"\imagem_drenagem_coletiva.jpg"
        texto3_drenagem_coletiva = "\t" + df["TEXTO2"][9]

        subtitulo_vapor_preso = "2.3.2." + "\t" + df['TÍTULO'][10]
        texto_vapor_preso = "\t" + df["TEXTO"][10]
        texto2_vapor_preso = "\t" + df["Figuras"][10]
        texto3_vapor_preso = "\t" + df["TEXTO1"][0]
        texto4_vapor_preso = "\t" + df["TEXTO2"][10]
        texto5_vapor_preso = "\t" + df["TEXTO3"][10]
        texto6_vapor_preso = "\t" + df["TEXTO4"][10]
        texto7_vapor_preso = "\t" + df["TEXTO5"][10]
        imagem_vapor_preso = self.rota_imagens + r"\imagem_vapor_preso.jpg"
        texto3_vapor_preso = "\t" + df["TEXTO2"][10]

        subtitulo_vapor_preso_cilindros = "2.3.3." + "\t" + df['TÍTULO'][11]
        texto_vapor_preso_cilindros = "\t" + df["TEXTO"][11]
        imagem_vapor_preso_cilindros = self.rota_imagens + r"\imagem_vapor_preso_cilindros.jpg"
        texto2_vapor_preso_cilindros = "\t" + df["TEXTO1"][11]
        imagem2_vapor_preso_cilindros = self.rota_imagens + r"\imagem2_vapor_preso_cilindros.jpg"
        texto3_vapor_preso_cilindros = "\t" + df["TEXTO3"][11]
        imagem3_vapor_preso_cilindros = self.rota_imagens + r"\imagem3_vapor_preso_cilindros.jpg"
        texto4_vapor_preso_cilindros = "\t" + df["TEXTO5"][11]
        texto5_vapor_preso_cilindros = "\t" + df["TEXTO6"][11]
        imagem4_vapor_preso_cilindros = self.rota_imagens + r"\imagem4_vapor_preso_cilindros.jpg"

        subtitulo_isolamento = "2.4.1." + "\t" + df['TÍTULO'][12]
        texto_isolamento = "\t" + df["TEXTO"][12]
        tabela_isolamento = self.rota_imagens + r"\tabela_isolamento.jpg"
        tabela2_isolamento = self.rota_imagens + r"\tabela2_isolamento.jpg"

        subtitulo_boia_termo = "2.4.2." + "\t" + df['TÍTULO'][13]
        texto_boia_termo = "\t" + df["TEXTO"][13]
        grafico_boia_termo = self.rota_imagens + r"\grafico_boia_termo.jpg"
        texto2_boia_termo = "\t" + df["TEXTO1"][13]

        subtitulo_vaz_ext = "2.4.3." + "\t" + df['TÍTULO'][14]
        texto_vaz_ext = "\t" + df["TEXTO"][14]
        tabela_vaz_ext = self.rota_imagens + r"\tabela_vaz_ext.jpg"

        subtitulo_bypass = "2.4.4." + "\t" + df['TÍTULO'][15]
        texto_bypass = "\t" + df["TEXTO"][15]
        imagem_bypass = self.rota_imagens + r"\imagem_bypass.jpg"

        # ADICIONANDO O TITULO NO DOCUMENTO
        document.add_paragraph('', style='Titulo')
        document.add_paragraph('', style='Titulo')
        document.add_paragraph('', style='Titulo')
        titulo_add = document.add_paragraph(titulo, style='Titulo')
        titulo_add.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        if self.var_introducao.get():
            # ADICIONANDO INTRODUCAO
            document.add_page_break()
            document.add_paragraph(subtitulo_introducao, style='Subtitulo')
            document.add_paragraph()
            paragrafo_introducao = document.add_paragraph(texto_introducao, style='texto')
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            document.add_paragraph()
            paragrafo_introducao = document.add_paragraph(texto2_introducao, style='texto')
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            document.add_paragraph()
            paragrafo_introducao.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        if self.var_melhorias.get():
            # ADICIONANDO MELHORIAS
            document.add_page_break()
            document.add_paragraph(subtitulo_melhorias, style='Subtitulo')
            document.add_paragraph()
            paragrafo_melhorias = document.add_paragraph(texto_melhorias, style='texto')
            paragrafo_melhorias.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            document.add_picture(imagem_melhorias, width=Cm(15.13))
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if self.var_geracao_de_vapor.get():
            # ADICIONANDO GERACAO DE VAPOR
            document.add_page_break()
            document.add_paragraph(subtitulo_geracao, style='Subtitulo')
            document.add_paragraph()
            document.add_paragraph(texto_geracao, style='texto')
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            document.add_paragraph()
            document.add_paragraph(texto2_geracao, style='texto')
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            document.add_picture(tabela_geracao)
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            document.add_paragraph(legenda_geracao, style="texto")
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if self.var_distribuicao_de_vapor.get():
            # ADICIONANDO DISTRIBUICAO
            document.add_page_break()
            document.add_paragraph(subtitulo_distribuicao, style='Subtitulo')
            document.add_paragraph()
            document.add_paragraph(texto_distribuicao, style="texto")
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            document.add_paragraph(texto2_distribuicao, style='texto')
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            document.add_paragraph()
            document.add_picture(imagem_distribuicao)
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            document.add_paragraph(texto3_distribuicao, style='texto')
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            last_paragraph = document.paragraphs[-1]
            document.add_paragraph(texto4_distribuicao, style='texto')
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            document.add_paragraph(texto5_distribuicao, style='texto')
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            document.add_paragraph(texto6_distribuicao, style='texto')
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            document.add_paragraph(texto7_distribuicao, style='texto')
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            if len(self.fotos_distribuicao_de_vapor) >= 1:
                document.add_page_break()
                for self.foto in self.fotos_distribuicao_de_vapor:
                    document.add_paragraph("•\tTAG - NOME DO LOCAL", style='texto')
                    document.add_paragraph()
                    document.add_picture(self.foto, height=Cm(6))

        if self.var_boas_praticas.get():
            # ADICIONANDO BOAS PRATICAS
            document.add_page_break()
            document.add_paragraph(subtitulo_boaspraticas, style='Subtitulo')
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            document.add_paragraph()
            document.add_paragraph(texto_boaspraticas, style="texto")
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            document.add_paragraph(texto2_boaspraticas, style='texto')
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            document.add_paragraph(texto3_boaspraticas, style='texto')
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            document.add_paragraph(texto4_boaspraticas, style='texto')
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            document.add_paragraph(texto5_boaspraticas, style='texto')
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            document.add_picture(imagem_boaspraticas)
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            document.add_paragraph(texto6_boaspraticas, style="texto")
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            document.add_picture(imagem2_boaspraticas)
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if len(self.fotos_boas_praticas) >= 1:
                document.add_page_break()
                for self.foto in self.fotos_boas_praticas:
                    document.add_paragraph("•\tTAG - NOME DO LOCAL", style='texto')
                    document.add_paragraph()
                    document.add_picture(self.foto, height=Cm(6))

        if self.var_adeq_in_out.get():
            # ADICIONANDO NECESSIDADE DE DRENAGEM ANTES DE REDUCAO / CONTROLE
            document.add_page_break()
            document.add_paragraph(subtitulo_antesred, style='Subtitulo')
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            document.add_paragraph()
            document.add_paragraph(texto_antesred, style="texto")
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            document.add_picture(imagem_antesred)
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # document.add_picture(imagem2_antesred)
            # last_paragraph = document.paragraphs[-1]
            # last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if len(self.fotos_adeq_in_out) >= 1:
                document.add_page_break()
                for self.foto in self.fotos_adeq_linhas_ext:
                    document.add_paragraph("•\tTAG - NOME DO LOCAL", style='texto')
                    document.add_paragraph()
                    document.add_picture(self.foto, height=Cm(6))

        if self.var_adeq_linhas_ext.get():
            # ADICIONANDO ADEQUACAO DE LINHAS EXISTENTES
            document.add_page_break()
            document.add_paragraph(subtitulo_adequacao, style='Subtitulo')
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            document.add_paragraph()
            document.add_paragraph(texto_adequacao, style="texto")
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            document.add_paragraph()
            document.add_paragraph()
            document.add_picture(imagem_adequacao)
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            document.add_picture(imagem2_adequacao)
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            document.add_picture(tabela_adequacao)
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            document.add_paragraph(problema_adequacao, style='texto')
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            if len(self.fotos_adeq_linhas_ext) >= 1:
                for self.foto in self.fotos_adeq_linhas_ext:
                    document.add_paragraph("•\tTAG - NOME DO LOCAL", style='texto')
                    document.add_paragraph()
                    document.add_picture(self.foto, height=Cm(6))

        if self.var_desaer_rede.get():
            # ADICIONANDO DESAERACAO
            document.add_page_break()
            document.add_paragraph(subtitulo_desaeracao, style='Subtitulo')
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            document.add_paragraph()
            document.add_paragraph(texto_desaeracao, style="texto")
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            document.add_picture(imagem_desaeracao)
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            document.add_picture(imagem2_desaeracao)
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            document.add_paragraph(nota_desaeracao, style='texto')
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            if len(self.fotos_desaeracao) >= 1:
                document.add_page_break()
                for self.foto in self.fotos_desaeracao:
                    document.add_paragraph("•\tTAG - NOME DO LOCAL", style='texto')
                    document.add_paragraph()
                    document.add_picture(self.foto, height=Cm(6))

        if self.var_dren_coletiva.get():
            # ADICIONANDO DRENAGEM COLETIVA
            document.add_page_break()
            document.add_paragraph(subtitulo_drenagem_coletiva, style='Subtitulo')
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            document.add_paragraph()
            document.add_paragraph(texto_drenagem_coletiva, style="texto")
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            document.add_paragraph(texto2_drenagem_coletiva, style="texto")
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            document.add_picture(imagem_drenagem_coletiva)
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            document.add_paragraph(texto3_drenagem_coletiva, style="texto")
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            if len(self.fotos_drenagem_coletiva) >= 1:
                document.add_page_break()
                for self.foto in self.fotos_drenagem_coletiva:
                    document.add_paragraph("•\tTAG - NOME DO LOCAL", style='texto')
                    document.add_paragraph()
                    document.add_picture(self.foto, height=Cm(6))

        if self.var_vapor_preso_teoria.get():
            # ADICIONANDO VAPOR PRESO TEORIA
            document.add_page_break()
            document.add_paragraph(subtitulo_vapor_preso, style='Subtitulo')
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            document.add_paragraph()
            document.add_paragraph(texto_vapor_preso, style="texto")
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            document.add_paragraph()
            document.add_paragraph(texto2_vapor_preso, style="texto")
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            document.add_paragraph()
            document.add_paragraph(texto3_vapor_preso, style="texto")
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            document.add_paragraph()
            document.add_picture(imagem_vapor_preso)
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            document.add_paragraph(texto4_vapor_preso, style="texto")
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            document.add_paragraph()
            document.add_paragraph(texto5_vapor_preso, style="texto")
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            document.add_paragraph()
            document.add_paragraph(texto6_vapor_preso, style="texto")
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            document.add_paragraph()
            document.add_paragraph(texto7_vapor_preso, style="texto")
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            document.add_paragraph()
            if len(self.fotos_vapor_preso_teoria) >= 1:
                document.add_page_break()
                for self.foto in self.fotos_vapor_preso_teoria:
                    document.add_paragraph("•\tTAG - NOME DO LOCAL", style='texto')
                    document.add_paragraph()
                    document.add_picture(self.foto, height=Cm(6))

        if self.var_vapor_preso_cilindros.get():
            # ADICIONANDO VAPOR PRESO CILINDROS
            document.add_page_break()
            document.add_paragraph(subtitulo_vapor_preso_cilindros, style='Subtitulo')
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            document.add_paragraph()
            document.add_paragraph(texto_vapor_preso_cilindros, style="texto")
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            document.add_picture(imagem_vapor_preso_cilindros)
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            document.add_paragraph(texto2_vapor_preso_cilindros, style="texto")
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            document.add_picture(imagem2_vapor_preso_cilindros)
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            document.add_paragraph(texto3_vapor_preso_cilindros, style="texto")
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            document.add_picture(imagem3_vapor_preso_cilindros)
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            document.add_paragraph(texto4_vapor_preso_cilindros, style="texto")
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            document.add_paragraph(texto5_vapor_preso_cilindros, style="texto")
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            document.add_picture(imagem4_vapor_preso_cilindros)
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if len(self.fotos_vapor_preso_cilindros) >= 1:
                document.add_page_break()
                for self.foto in self.fotos_vapor_preso_cilindros:
                    document.add_paragraph("•\tTAG - NOME DO LOCAL", style='texto')
                    document.add_paragraph()
                    document.add_picture(self.foto, height=Cm(6))

        if self.var_tubos_isolamento.get():
            # ADICIONANDO ISOLAMENTO
            document.add_page_break()
            document.add_paragraph(subtitulo_isolamento, style='Subtitulo')
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            document.add_paragraph()
            document.add_paragraph(texto_isolamento, style="texto")
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            document.add_picture(tabela_isolamento)
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            document.add_picture(tabela2_isolamento)
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if len(self.fotos_tubos_isolamento) >= 1:
                document.add_page_break()
                for self.foto in self.fotos_tubos_isolamento:
                    document.add_paragraph("•\tTAG - NOME DO LOCAL", style='texto')
                    document.add_paragraph()
                    document.add_picture(self.foto, height=Cm(6))

        if self.var_purgadores_termo_x_boia.get():
            # ADICIONANDO BOIA TERMO
            document.add_page_break()
            document.add_paragraph(subtitulo_boia_termo, style='Subtitulo')
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            document.add_paragraph()
            document.add_paragraph(texto_boia_termo, style="texto")
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            document.add_picture(grafico_boia_termo, width=Cm(15))
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            document.add_paragraph(texto2_boia_termo, style='texto')
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        if self.var_locais_vazamentos_ext.get():
            # ADICIONANDO VAZAMENTO EXTERNO
            document.add_page_break()
            document.add_paragraph(subtitulo_vaz_ext, style='Subtitulo')
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            document.add_paragraph()
            document.add_paragraph(texto_vaz_ext, style="texto")
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            document.add_picture(tabela_vaz_ext, width=Cm(10.13))
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if len(self.fotos_locais_vaz_ext) >= 1:
                document.add_page_break()
                for self.foto in self.fotos_locais_vaz_ext:
                    document.add_paragraph("•\tTAG - NOME DO LOCAL", style='texto')
                    document.add_paragraph()
                    document.add_picture(self.foto, height=Cm(6))

        if self.var_valvulas_bypass.get():
            # ADICIONANDO BYPASS
            document.add_page_break()
            document.add_paragraph(subtitulo_bypass, style='Subtitulo')
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            document.add_paragraph()
            document.add_paragraph(texto_bypass, style="texto")
            last_paragraph = document.paragraphs[-1]
            last_paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            document.add_picture(imagem_bypass, height=Cm(3))
            last_paragraph = document.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if len(self.fotos_valvulas_bypass) >= 1:
                document.add_page_break()
                for self.foto in self.fotos_valvulas_bypass:
                    document.add_paragraph("•\tTAG - NOME DO LOCAL", style='texto')
                    document.add_picture(self.foto, height=Cm(23.61))

        document.save(self.save_name)
        messagebox.showinfo("Gerador de Relatório - TMS", f"Documento {self.save_name} gerado com sucesso!")

    def check_status_distribuicao(self):
        if self.var_distribuicao_de_vapor.get():
            self.caminho_distribuicao = filedialog.askopenfilenames()
            self.fotos_distribuicao_de_vapor = list(self.caminho_distribuicao)
        return self.fotos_distribuicao_de_vapor

    def check_status_boas_praticas(self):
        if self.var_boas_praticas.get():
            self.caminho = filedialog.askopenfilenames()
            self.fotos_boas_praticas = list(self.caminho)
        return self.fotos_boas_praticas

    def check_status_antes_red(self):
        if self.var_adeq_in_out.get():
            self.caminho = filedialog.askopenfilenames()
            self.fotos_adeq_in_out = list(self.caminho)
        return self.fotos_adeq_in_out

    def check_status_adeq_existente(self):
        if self.var_adeq_linhas_ext.get():
            self.caminho = filedialog.askopenfilenames()
            self.fotos_adeq_linhas_ext = list(self.caminho)
        return self.fotos_adeq_linhas_ext

    def check_status_desaeracao(self):
        if self.var_desaer_rede.get():
            self.caminho = filedialog.askopenfilenames()
            self.fotos_desaeracao = list(self.caminho)
        return self.fotos_desaeracao

    def check_status_drenagem_coletiva(self):
        if self.var_dren_coletiva.get():
            self.caminho = filedialog.askopenfilenames()
            self.fotos_drenagem_coletiva = list(self.caminho)
        return self.fotos_drenagem_coletiva

    def check_status_vapor_preso(self):
        if self.var_vapor_preso_teoria.get():
            self.caminho = filedialog.askopenfilenames()
            self.fotos_vapor_preso_teoria = list(self.caminho)
        return self.fotos_vapor_preso_teoria

    def check_status_vapor_preso_cilindros(self):
        if self.var_vapor_preso_cilindros.get():
            self.caminho = filedialog.askopenfilenames()
            self.fotos_vapor_preso_cilindros = list(self.caminho)
        return self.fotos_vapor_preso_cilindros

    def check_status_tubos_isolamento(self):
        if self.var_tubos_isolamento.get():
            self.caminho = filedialog.askopenfilenames()
            self.fotos_tubos_isolamento = list(self.caminho)
        return self.fotos_tubos_isolamento

    def check_status_locais_vaz_ext(self):
        if self.var_locais_vazamentos_ext.get():
            self.caminho = filedialog.askopenfilenames()
            self.fotos_locais_vaz_ext = list(self.caminho)
        return self.fotos_locais_vaz_ext

    def check_status_valvulas_bypass(self):
        if self.var_valvulas_bypass.get():
            self.caminho = filedialog.askopenfilenames()
            self.fotos_valvulas_bypass = list(self.caminho)
        return self.fotos_valvulas_bypass


if __name__ == "__main__":
    app = tk.Tk()
    app.title("Gerador de Relatório - TMS")
    MainApp(app)
    app.mainloop()
